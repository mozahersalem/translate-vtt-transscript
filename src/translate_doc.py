from docx import Document
from google.cloud import translate_v2 as translate
import os
import concurrent.futures

def translate_text(text, target_language):
    """Translates a single text string using Google Cloud Translate API."""
    try:
        translate_client = translate.Client()  # Initialize inside the function
        translation = translate_client.translate(text, target_language=target_language)
        return translation.get("translatedText", text)
    except Exception as e:
        print(f"⚠️ ERROR: Google Translate API failed for '{text[:50]}...' → {e}")
        return text  # Return original text on failure

def translate_doc(input_docx_path, output_docx_path, target_language='en'):
    """Translates a DOCX file using Google Cloud Translate API."""

    if "GOOGLE_APPLICATION_CREDENTIALS" not in os.environ:
        raise EnvironmentError("❌ ERROR: Google Cloud authentication is missing! Run 'gcloud auth application-default login'.")

    try:
        # ✅ Load the original DOCX file
        doc = Document(input_docx_path)
        translated_doc = Document()

        print(f"📖 Translating document: {input_docx_path} → {target_language}")
        print(f"📝 Total paragraphs: {len(doc.paragraphs)}")

        paragraphs = [paragraph for paragraph in doc.paragraphs if paragraph.text.strip()]
        translated_texts = [None] * len(paragraphs)

        with concurrent.futures.ThreadPoolExecutor(max_workers=5) as executor:
            future_to_index = {
                executor.submit(translate_text, paragraph.text.strip(), target_language): i
                for i, paragraph in enumerate(paragraphs)
            }

            for future in concurrent.futures.as_completed(future_to_index):
                index = future_to_index[future]
                try:
                    translated_texts[index] = future.result()
                    # print(f"✅ Translated paragraph {index+1}/{len(paragraphs)}: {translated_texts[index][:50]}...")
                except Exception as e:
                    print(f"⚠️ ERROR: {e}")
                    translated_texts[index] = paragraphs[index].text  # Keep original text on failure

        for i, paragraph in enumerate(doc.paragraphs):
            new_paragraph = translated_doc.add_paragraph()
            translated_text = translated_texts[i]
            for run in paragraph.runs:
                new_run = new_paragraph.add_run(translated_text)
                new_run.bold = run.bold
                new_run.italic = run.italic
                new_run.underline = run.underline

        translated_doc.save(output_docx_path)
        print(f"✅ Translated document saved as: {output_docx_path}")

    except Exception as e:
        print(f"❌ ERROR: {e}")

# Example Usage
# if __name__ == "__main__":
#     input_file = "input.docx"
#     output_file = "translated.docx"
#     target_lang = "es"  # Change language

#     translate_doc(input_file, output_file, target_lang)